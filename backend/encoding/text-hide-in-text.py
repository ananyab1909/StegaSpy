from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import docx
import glob

def get_full_path(file_name):
    for root, dirs, files in os.walk(os.getcwd()):
        for file in glob.glob(os.path.join(root, file_name)):
            return os.path.abspath(file)
    return "File not found"

def hide_secret_message(text, secret_message):
    doc = docx.Document()
    
    for i, char in enumerate(text):
        para = doc.add_paragraph()
        
        if i < len(secret_message):
            run = para.add_run(secret_message[i])
            run.font.size = docx.shared.Pt(10) 
        else:
            run = para.add_run(char)
            run.font.size = docx.shared.Pt(12) 
    
    doc.save('output.docx')

app=Flask(__name__)
CORS(app)

@app.route('/decode_text', methods=['POST'])
def decode_text():
    try:
        if not request.is_json:
          return jsonify({'status': 400, 'message': 'Invalid request'}), 400
        
        data = request.json
        if 'coverText' not in data or 'secretMessage' not in data:
            return jsonify({'status': 400, 'message': 'Invalid request'}), 400
        
        cover = data['coverText']
        message = data['secretMessage']
        
        try:
            hide_secret_message(cover,message)
            return jsonify({'status': 200, 'document': 'output.docx', 'message': 'Success'}), 200
        except Exception as e:
            print(f"Error processing request: {e}")
            return jsonify({'status': 500, 'message': 'Internal Server Error'}), 500

        
    except Exception as e:
        import traceback
        error_message = traceback.format_exc()
        print(f"Error: {error_message}")
        return jsonify({'status': 500, 'message': 'Internal Server Error'}), 500
    
if __name__ == '__main__':
    app.run(debug=True)